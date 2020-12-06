import docx
# pip install python-docx 

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def saveDoc(filename, text):
    newDoc = docx.Document()
    newDoc.add_paragraph(text)
    newDoc.save(filename)

def removePunctuation(text):
    punctuations = ".,?!#:;"
    cleanText = ""
    for char in text:
        if char not in punctuations:
            cleanText += char
    return cleanText.lower()

docText = getText("Corpus/Clément 2.docx")
cleanText = removePunctuation(docText)
saveDoc("Corpus/Clément 2NP.docx", cleanText)

